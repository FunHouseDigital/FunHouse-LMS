"""Property-based tests for the Extract lesson ``.docx`` path (Tasks 7.2-7.4).

Implements design Properties 9, 10 and 11. Each test builds **synthetic
``.docx`` fixtures** with ``python-docx`` in a temporary directory and drives the
deterministic text parser -- there is no network and no LLM. Each property runs a
minimum of 100 Hypothesis iterations, per the design's Testing Strategy.
"""

from __future__ import annotations

import tempfile
from datetime import datetime, timezone
from pathlib import Path
from unittest import mock

import pytest
from docx import Document
from hypothesis import given, settings
from hypothesis import strategies as st

from funhouse_pipeline.collect import HandlerTarget, RoutedFile
from funhouse_pipeline.extract import (
    ALLOWED_METRIC_TYPES,
    extract_lessons,
)

pytestmark = pytest.mark.property

_SETTINGS = settings(max_examples=100, deadline=None)

_ALLOWED_SET = set(ALLOWED_METRIC_TYPES)

_FIXED = datetime(2024, 1, 1, tzinfo=timezone.utc)

# Safe, human-ish text: letters + spaces, trimmed, non-empty.
_words = (
    st.text(
        alphabet="abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ ",
        min_size=1,
        max_size=18,
    )
    .map(lambda s: " ".join(s.split()))
    .filter(lambda s: len(s) >= 1)
)

# Raw metric labels: a mix of recognized (canonical + keyword) and unrecognized.
_metric_labels = st.sampled_from(
    [
        # canonical / allowed
        "typing_wpm",
        "typing_accuracy",
        "homework_done",
        "quiz_score",
        "observation",
        # keyword variants that MUST normalize into the allowed set
        "WPM",
        "Words Per Minute",
        "Accuracy",
        "Homework",
        "Quiz",
        "Observation",
        "Notes",
        # unrecognized -> must NOT produce a record
        "height",
        "age",
        "favourite colour",
        "",
    ]
)


def _routed(path: Path) -> RoutedFile:
    return RoutedFile(
        path=path,
        subfolder="lessons",
        source_type="lesson documents",
        handler=HandlerTarget.DOCX_TEXT_PARSER,
    )


def _write_single_lesson_doc(path: Path, *, title: str, body: list[str]) -> None:
    doc = Document()
    doc.add_paragraph(f"LESSON: {title}")
    for line in body:
        doc.add_paragraph(line)
    doc.save(str(path))


# Feature: phase0-data-foundation, Property 9: Lesson `.docx` files are parsed as
# text without OCR or LLM image calls. For any `.docx` lesson file, extraction
# produces records via the text parser and issues no image-OCR or LLM
# image-extraction call.
# Validates: Requirements 5.1
@_SETTINGS
@given(
    title=_words,
    topic=_words,
    phenomenon=_words,
    content_lines=st.lists(_words, min_size=0, max_size=5),
)
def test_property_9_docx_parsed_as_text_without_llm(
    title, topic, phenomenon, content_lines
):
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "lesson.docx"
        body = [f"TOPIC: {topic}", f"PHENOMENON: {phenomenon}", *content_lines]
        _write_single_lesson_doc(path, title=title, body=body)

        # Any LLM call whatsoever would fail this test: patch the single model
        # entry point and assert it is never touched by the text path (Req 5.1).
        with mock.patch("funhouse_pipeline.llm.llm_generate") as llm_spy:
            records = extract_lessons([_routed(path)], now=lambda: _FIXED)

        assert llm_spy.call_count == 0
        # The text parser produced records, all tagged with the docx provider.
        assert records, "text parser must produce at least the one lessons record"
        assert all(r.provider == "docx-parser" for r in records)
        lessons = [r for r in records if r.target_table == "lessons"]
        assert len(lessons) == 1


def _write_multi_lesson_doc(path: Path, titles: list[str]) -> None:
    doc = Document()
    for i, title in enumerate(titles):
        doc.add_paragraph(f"LESSON: {title}")
        doc.add_paragraph(f"TOPIC: topic {i}")
        doc.add_paragraph(f"some content for lesson {i}")
    doc.save(str(path))


# Feature: phase0-data-foundation, Property 10: One lessons record per lesson.
# For any lesson document containing N lessons, extraction produces exactly N
# `lessons` records.
# Validates: Requirements 5.2, 10.1
@_SETTINGS
@given(titles=st.lists(_words, min_size=1, max_size=6))
def test_property_10_one_lessons_record_per_lesson(titles):
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "multi.docx"
        _write_multi_lesson_doc(path, titles)

        records = extract_lessons([_routed(path)], now=lambda: _FIXED)

        lessons = [r for r in records if r.target_table == "lessons"]
        # Exactly N lessons for N `LESSON:` markers (documented delimiting rule).
        assert len(lessons) == len(titles)


# --------------------------------------------------------------------------- #
# Property 11 generators: build docs whose metrics arrive via both METRIC: lines
# and a metric table, mixing recognized and unrecognized labels.
# --------------------------------------------------------------------------- #

_metric_value = st.one_of(
    st.integers(min_value=0, max_value=120).map(str),
    _words,
)


def _write_metrics_doc(path: Path, metric_lines, table_headers, table_rows) -> None:
    doc = Document()
    doc.add_paragraph("LESSON: Metrics Lesson")
    for player, label, value in metric_lines:
        doc.add_paragraph(f"METRIC: {player} | {label} | {value}")
    if table_headers:
        table = doc.add_table(rows=1 + len(table_rows), cols=1 + len(table_headers))
        header_cells = table.rows[0].cells
        header_cells[0].text = "Student"
        for j, header in enumerate(table_headers):
            header_cells[j + 1].text = header
        for i, row_values in enumerate(table_rows):
            cells = table.rows[i + 1].cells
            cells[0].text = f"Player {i}"
            for j, val in enumerate(row_values):
                cells[j + 1].text = str(val)
    doc.save(str(path))


# Feature: phase0-data-foundation, Property 11: Embedded metrics use only allowed
# metric types. For any lesson document with embedded learning measurements, each
# produced `student_metrics` record has a `metric_type` from the allowed set.
# Validates: Requirements 5.3
@_SETTINGS
@given(
    metric_lines=st.lists(
        st.tuples(_words, _metric_labels, _metric_value), min_size=0, max_size=5
    ),
    table_headers=st.lists(_metric_labels, min_size=0, max_size=4),
    n_rows=st.integers(min_value=0, max_value=3),
    cell_values=st.lists(_metric_value, min_size=0, max_size=12),
)
def test_property_11_embedded_metrics_use_only_allowed_types(
    metric_lines, table_headers, n_rows, cell_values
):
    # Build rectangular table rows from the flat cell_values pool.
    table_rows = []
    cursor = 0
    width = len(table_headers)
    for _ in range(n_rows if width else 0):
        row = []
        for _ in range(width):
            row.append(cell_values[cursor % len(cell_values)] if cell_values else "1")
            cursor += 1
        table_rows.append(row)

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "metrics.docx"
        _write_metrics_doc(path, metric_lines, table_headers, table_rows)

        records = extract_lessons([_routed(path)], now=lambda: _FIXED)

        metrics = [r for r in records if r.target_table == "student_metrics"]
        for record in metrics:
            assert record.payload["metric_type"] in _ALLOWED_SET
