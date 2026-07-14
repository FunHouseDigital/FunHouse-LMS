"""Unit / example tests for the Extract lesson ``.docx`` path (Task 7.1).

Concrete, readable scenarios that complement the property tests: one file = one
lesson by default, ``LESSON:`` delimiting, topic/phenomenon tagging, embedded
metric recognition (both ``METRIC:`` lines and metric tables), the deterministic
envelope, and the guarantee that the text path imports no LLM/provider SDK.
Pure local work with synthetic ``.docx`` fixtures -- no network.
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document

from funhouse_pipeline.collect import HandlerTarget, RoutedFile
from funhouse_pipeline.extract import (
    ALLOWED_METRIC_TYPES,
    extract_lessons,
    parse_lesson_document,
)

_FIXED = datetime(2024, 5, 6, 7, 8, 9, tzinfo=timezone.utc)


def _routed(path: Path) -> RoutedFile:
    return RoutedFile(
        path=path,
        subfolder="lessons",
        source_type="lesson documents",
        handler=HandlerTarget.DOCX_TEXT_PARSER,
    )


def test_single_lesson_default_one_file_one_lesson(tmp_path):
    path = tmp_path / "intro_to_typing.docx"
    doc = Document()
    doc.add_paragraph("TOPIC: Keyboard basics")
    doc.add_paragraph("PHENOMENON: Muscle memory")
    doc.add_paragraph("Today we learn the home row.")
    doc.save(str(path))

    records = extract_lessons([_routed(path)], now=lambda: _FIXED)

    lessons = [r for r in records if r.target_table == "lessons"]
    assert len(lessons) == 1
    lesson = lessons[0]
    # No LESSON: marker -> title defaults to the file stem.
    assert lesson.payload["title"] == "intro_to_typing"
    assert lesson.payload["topic"] == "Keyboard basics"
    assert lesson.payload["phenomenon"] == "Muscle memory"
    assert "home row" in lesson.payload["content"]
    assert lesson.payload["source_file"] == str(path)
    # Deterministic envelope.
    assert lesson.provider == "docx-parser"
    assert lesson.confidence_score == 1.0
    assert lesson.extracted_at == _FIXED
    assert lesson.record_id


def test_multiple_lessons_delimited_by_marker(tmp_path):
    path = tmp_path / "unit.docx"
    doc = Document()
    doc.add_paragraph("LESSON: Fractions")
    doc.add_paragraph("TOPIC: Numbers")
    doc.add_paragraph("LESSON: Decimals")
    doc.add_paragraph("TOPIC: Numbers")
    doc.add_paragraph("LESSON: Geometry")
    doc.save(str(path))

    records = extract_lessons([_routed(path)], now=lambda: _FIXED)
    lessons = [r for r in records if r.target_table == "lessons"]
    assert [lsn.payload["title"] for lsn in lessons] == [
        "Fractions",
        "Decimals",
        "Geometry",
    ]


def test_metric_line_recognition_and_keyword_normalization(tmp_path):
    path = tmp_path / "lesson.docx"
    doc = Document()
    doc.add_paragraph("LESSON: Typing")
    doc.add_paragraph("METRIC: Thabo | WPM | 42")
    doc.add_paragraph("METRIC: Thabo | Accuracy | 95%")
    doc.add_paragraph("METRIC: Naledi | quiz_score | 8")
    # Unrecognized metric type -> not emitted.
    doc.add_paragraph("METRIC: Naledi | height | 150")
    doc.save(str(path))

    records = extract_lessons([_routed(path)], now=lambda: _FIXED)
    metrics = [r for r in records if r.target_table == "student_metrics"]
    types = sorted(m.payload["metric_type"] for m in metrics)
    assert types == ["quiz_score", "typing_accuracy", "typing_wpm"]
    assert all(m.payload["metric_type"] in ALLOWED_METRIC_TYPES for m in metrics)
    # Metrics are attached to their enclosing lesson.
    assert all(m.payload["lesson_title"] == "Typing" for m in metrics)


def test_metric_table_recognition(tmp_path):
    path = tmp_path / "typingbird.docx"
    doc = Document()
    doc.add_paragraph("LESSON: TypingBird Session")
    table = doc.add_table(rows=3, cols=4)
    headers = table.rows[0].cells
    headers[0].text = "Student"
    headers[1].text = "WPM"
    headers[2].text = "Accuracy"
    headers[3].text = "Age"  # unrecognized column -> ignored
    r1 = table.rows[1].cells
    r1[0].text = "Thabo"
    r1[1].text = "40"
    r1[2].text = "92"
    r1[3].text = "11"
    r2 = table.rows[2].cells
    r2[0].text = "Naledi"
    r2[1].text = "55"
    r2[2].text = ""  # empty cell -> no record
    r2[3].text = "10"
    doc.save(str(path))

    records = extract_lessons([_routed(path)], now=lambda: _FIXED)
    metrics = [r for r in records if r.target_table == "student_metrics"]
    got = sorted(
        (m.payload["player_name"], m.payload["metric_type"], m.payload["value"])
        for m in metrics
    )
    assert got == [
        ("Naledi", "typing_wpm", "55"),
        ("Thabo", "typing_accuracy", "92"),
        ("Thabo", "typing_wpm", "40"),
    ]


def test_non_docx_and_image_routed_files_are_ignored(tmp_path):
    docx = tmp_path / "lesson.docx"
    doc = Document()
    doc.add_paragraph("LESSON: X")
    doc.save(str(docx))

    image = RoutedFile(
        path=tmp_path / "card.png",
        subfolder="cards",
        source_type="membership/pay cards",
        handler=HandlerTarget.IMAGE_EXTRACT,
    )

    records = extract_lessons([image, _routed(docx)], now=lambda: _FIXED)
    lessons = [r for r in records if r.target_table == "lessons"]
    assert len(lessons) == 1
    assert lessons[0].payload["title"] == "X"


def test_direct_path_accepted(tmp_path):
    path = tmp_path / "direct.docx"
    doc = Document()
    doc.add_paragraph("LESSON: Direct")
    doc.save(str(path))

    records = extract_lessons([path], now=lambda: _FIXED)
    assert [r.payload["title"] for r in records if r.target_table == "lessons"] == [
        "Direct"
    ]


def test_lessons_module_imports_no_llm_or_provider_sdk():
    from funhouse_pipeline.extract import lessons as lessons_mod

    forbidden = {
        "boto3",
        "anthropic",
        "llm_generate",
        "BedrockBatchProvider",
        "AnthropicProvider",
    }
    leaked = forbidden.intersection(vars(lessons_mod))
    assert not leaked, f"lesson text path must not reference an LLM/SDK; found {leaked}"


def test_empty_document_still_yields_one_lesson(tmp_path):
    path = tmp_path / "empty.docx"
    Document().save(str(path))

    lessons, source = parse_lesson_document(path)
    assert len(lessons) == 1
    assert source == str(path)
