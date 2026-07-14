"""Integration tests for the end-to-end command (Task 15.1).

These are concrete, focused examples (not property tests) that exercise the
whole pipeline -- **Collect -> Extract -> Validate -> Load -> Archive** -- across
its real component seams, with only the *external* services faked:

* a deterministic ``llm_generate`` stub for the image Extract path (no Bedrock),
* a **moto**-backed S3 for Archive (no live S3),
* an **ephemeral PostgreSQL schema** for Load (created per test, dropped after).

They verify:

1. A fresh ``Source_Folder`` (an image under ``photos/`` plus a lesson ``.docx``
   under ``lessons/``) processes fully through all five stages, landing rows in
   PostgreSQL and originals in S3 (Req 13.1, 13.4).
2. Setting ``LLM_PROVIDER=anthropic`` routes extraction through the Anthropic
   provider (stubbed client) with **no Extractor/pipeline code change** -- the
   provider swap is transparent (Req 6.2). The extracted records carry
   ``provider == "anthropic"`` and still load end to end.

DB-backed, so skipped automatically when no PostgreSQL server is reachable.
"""

from __future__ import annotations

import json
import uuid
from datetime import date

import pytest
from docx import Document

from funhouse_pipeline.config import Config, DatabaseConfig
from funhouse_pipeline.llm import ProviderRegistry, llm_generate
from funhouse_pipeline.llm.anthropic import AnthropicProvider
from funhouse_pipeline.llm.base import LLMResult, LLMResultItem
from funhouse_pipeline.orchestrator.pipeline import run_pipeline
from funhouse_pipeline.orchestrator.retry import RetryPolicy

moto = pytest.importorskip("moto")
from moto import mock_aws  # noqa: E402

pytestmark = pytest.mark.db

_REFERENCE_DATE = date(2025, 1, 1)


# --------------------------------------------------------------------------- #
# Fixture Source_Folder
# --------------------------------------------------------------------------- #


def _build_source_folder(root):
    """A small but representative Source_Folder: one image + one lesson .docx.

    * ``photos/scan.png`` -> image Extract path (fake LLM produces players + a
      payment referencing one of them).
    * ``lessons/typing.docx`` -> deterministic ``.docx`` text path (no LLM),
      producing one ``lessons`` record plus an embedded ``student_metrics`` row
      for a player the image also produced (so it resolves and loads).
    """
    photos = root / "photos"
    photos.mkdir(parents=True, exist_ok=True)
    (photos / "scan.png").write_bytes(b"\x89PNG\r\n\x1a\n fake image bytes")

    lessons = root / "lessons"
    lessons.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("LESSON: Home Row Basics")
    doc.add_paragraph("TOPIC: Keyboarding")
    doc.add_paragraph("PHENOMENON: Muscle memory")
    doc.add_paragraph("METRIC: Kabelo | WPM | 42")
    doc.save(str(lessons / "typing.docx"))
    return root


#: The records the fake model "extracts" from the single image file: two
#: players and a payment for one of them at a valid tier amount.
_IMAGE_PAYLOAD = {
    "records": [
        {
            "target_table": "players",
            "confidence": 0.95,
            "payload": {"first_name": "Kabelo", "last_name": "Test", "birth_date": "2015-06-01"},
        },
        {
            "target_table": "players",
            "confidence": 0.95,
            "payload": {"first_name": "Naledi", "last_name": "Test", "birth_date": "2016-07-02"},
        },
        {
            "target_table": "payments",
            "confidence": 0.95,
            "payload": {"player_name": "Kabelo Test", "amount": "R30"},
        },
    ]
}


def _fake_llm_generate(task, context, *, provider=None, options=None):
    """Deterministic image-path stub: every input record maps to _IMAGE_PAYLOAD."""
    items = [
        LLMResultItem(custom_id=rec["custom_id"], content=json.dumps(_IMAGE_PAYLOAD))
        for rec in context["records"]
    ]
    return LLMResult(task=task, provider=provider or "bedrock", items=tuple(items))


# --------------------------------------------------------------------------- #
# Ephemeral schema + connection factory
# --------------------------------------------------------------------------- #


def _connection_factory(dsn, schema):
    import psycopg

    def factory(_config: Config):
        # Autocommit so each loader per-record transaction persists (matches the
        # orchestrator's default connection behavior).
        conn = psycopg.connect(dsn, autocommit=True)
        conn.execute(f'SET search_path TO "{schema}", public')
        return conn

    return factory


def _count(conn, table):
    with conn.cursor() as cur:
        cur.execute(f"SELECT count(*) FROM {table}")
        return cur.fetchone()[0]


def _make_config():
    return Config(
        database=DatabaseConfig(host="127.0.0.1", dbname="funhouse_test", user="funhouse"),
        s3_bucket="funhouse-archive-integration",
        region="af-south-1",
        llm_provider="bedrock",
        confidence_threshold=0.7,
    )


# --------------------------------------------------------------------------- #
# 1. Full end-to-end run over a fresh Source_Folder (Req 13.1, 13.4)
# --------------------------------------------------------------------------- #


def test_fresh_source_folder_runs_all_five_stages(pg_dsn, tmp_path_factory):
    import psycopg

    schema = f"itg_{uuid.uuid4().hex}"
    admin = psycopg.connect(pg_dsn, autocommit=True)
    admin.execute(f'CREATE SCHEMA "{schema}"')

    src = _build_source_folder(tmp_path_factory.mktemp("src"))
    factory = _connection_factory(pg_dsn, schema)
    config = _make_config()
    policy = RetryPolicy(max_attempts=2, base_delay=0)

    try:
        with mock_aws():
            import boto3

            s3 = boto3.client("s3", region_name="af-south-1")
            s3.create_bucket(
                Bucket=config.s3_bucket,
                CreateBucketConfiguration={"LocationConstraint": "af-south-1"},
            )

            result = run_pipeline(
                config,
                src,
                state_dir=str(tmp_path_factory.mktemp("state")),
                llm_generate_fn=_fake_llm_generate,
                s3_client=s3,
                connection_factory=factory,
                retry_policy=policy,
                sleep=lambda _d: None,
                reference_date=_REFERENCE_DATE,
            )

            # All five stages ran end to end (Req 13.1).
            assert result.stages_run == ("collect", "extract", "validate", "load", "archive")

            # Collect saw both source files (Req 13.4: a new folder is processed
            # fully).
            assert result.summary["collected"] == 2

            # Extract produced records from BOTH paths: the image (players +
            # payment) and the .docx (lesson + embedded metric).
            tables = {r.target_table for r in result.records}
            assert {"players", "payments", "lessons", "student_metrics"} <= tables

            # Archive uploaded both originals under raw/ in S3.
            assert result.summary["archived"] == 2
            keys = sorted(o.key for o in result.archive_result.objects)
            assert keys == ["raw/lessons/typing.docx", "raw/photos/scan.png"]
            listed = s3.list_objects_v2(Bucket=config.s3_bucket).get("Contents", [])
            assert {o["Key"] for o in listed} == set(keys)

            # Load: the rows actually landed in PostgreSQL.
            read = factory(config)
            try:
                assert _count(read, "players") == 2          # deduped learners
                assert _count(read, "payments") == 1         # Kabelo's R30 payment
                assert _count(read, "lessons") == 1          # the one lesson
                assert _count(read, "student_metrics") == 1  # embedded WPM metric
                # Provenance: the lesson row points at the archived original.
                with read.cursor() as cur:
                    cur.execute("SELECT original_file_ref FROM lessons")
                    assert cur.fetchone()[0] == "raw/lessons/typing.docx"
            finally:
                read.close()

            # summary["loaded"] counts the non-player inserts (payment + lesson +
            # metric); the two deduped players are tracked in players.created.
            assert result.summary["loaded"] == 3
            assert len(result.load_result.players.created) == 2
    finally:
        admin.execute(f'DROP SCHEMA IF EXISTS "{schema}" CASCADE')
        admin.close()


# --------------------------------------------------------------------------- #
# 2. Provider swap: LLM_PROVIDER=anthropic routes through Anthropic (Req 6.2)
# --------------------------------------------------------------------------- #


class _FakeAnthropicMessages:
    """Stub of ``client.messages`` recording calls and returning fixed content."""

    def __init__(self, content: str) -> None:
        self._content = content
        self.calls: list[dict] = []

    def create(self, **kwargs):
        self.calls.append(kwargs)
        # Shape mirrors an Anthropic Messages API response (dict form is
        # accepted by AnthropicProvider._parse_response).
        return {"content": [{"type": "text", "text": self._content}], "stop_reason": "end_turn"}


class _FakeAnthropicClient:
    def __init__(self, content: str) -> None:
        self.messages = _FakeAnthropicMessages(content)


def test_provider_swap_to_anthropic_is_transparent(pg_dsn, tmp_path_factory):
    """Switching to the Anthropic provider requires no Extractor/pipeline code
    change: the same ``run_pipeline`` call routes extraction through the
    Anthropic provider (via ``LLM_PROVIDER=anthropic``) and still loads."""
    import psycopg

    schema = f"itg_{uuid.uuid4().hex}"
    admin = psycopg.connect(pg_dsn, autocommit=True)
    admin.execute(f'CREATE SCHEMA "{schema}"')

    src = _build_source_folder(tmp_path_factory.mktemp("src"))
    factory = _connection_factory(pg_dsn, schema)
    config = Config(
        database=DatabaseConfig(host="127.0.0.1", dbname="funhouse_test", user="funhouse"),
        s3_bucket="funhouse-archive-integration-anthropic",
        region="af-south-1",
        llm_provider="anthropic",  # <-- the only change vs. the bedrock run
        confidence_threshold=0.7,
    )
    policy = RetryPolicy(max_attempts=2, base_delay=0)

    # A registry whose "anthropic" provider is a real AnthropicProvider wired to
    # a stubbed client -- so the abstraction genuinely routes through the
    # Anthropic code path with no network.
    fake_client = _FakeAnthropicClient(json.dumps(_IMAGE_PAYLOAD))
    registry = ProviderRegistry()
    registry.register("anthropic", lambda _options: AnthropicProvider(client=fake_client))

    def anthropic_routing_llm(task, context, *, provider=None, options=None):
        # Identical to what the Extractor calls; provider is selected purely
        # from LLM_PROVIDER -- no Extractor code change.
        return llm_generate(
            task,
            context,
            registry=registry,
            env={"LLM_PROVIDER": provider or "anthropic"},
        )

    try:
        with mock_aws():
            import boto3

            s3 = boto3.client("s3", region_name="af-south-1")
            s3.create_bucket(
                Bucket=config.s3_bucket,
                CreateBucketConfiguration={"LocationConstraint": "af-south-1"},
            )

            result = run_pipeline(
                config,
                src,
                state_dir=str(tmp_path_factory.mktemp("state")),
                llm_generate_fn=anthropic_routing_llm,
                s3_client=s3,
                connection_factory=factory,
                retry_policy=policy,
                sleep=lambda _d: None,
                reference_date=_REFERENCE_DATE,
            )

            # The Anthropic provider actually served the extraction call.
            assert fake_client.messages.calls, "Anthropic client was never called"

            # Every image-extracted record is labeled with the anthropic provider.
            image_records = [r for r in result.records if r.source_file.endswith("scan.png")]
            assert image_records
            assert all(r.provider == "anthropic" for r in image_records)

            # And the swap is truly transparent: the pipeline still loaded rows.
            read = factory(config)
            try:
                assert _count(read, "players") == 2
                assert _count(read, "payments") == 1
            finally:
                read.close()
    finally:
        admin.execute(f'DROP SCHEMA IF EXISTS "{schema}" CASCADE')
        admin.close()
